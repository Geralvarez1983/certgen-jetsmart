#!/usr/bin/env python3
"""
CertGen - Generador de Certificados JetSMART
100% compatible con Windows - usa solo python-docx y openpyxl
"""

import os, shutil, zipfile, re, json, uuid, subprocess, sys
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
from docx import Document
from docx.shared import Inches
from copy import deepcopy
import openpyxl

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
# En Railway usar /data para persistencia, local usar carpeta del proyecto
DATA_ROOT   = os.environ.get('RAILWAY_DATA_DIR', BASE_DIR)
MODELOS_DIR = os.path.join(DATA_ROOT, 'modelos')
UPLOADS_DIR = os.path.join(DATA_ROOT, 'uploads')
DATA_FILE   = os.path.join(DATA_ROOT, 'data', 'modelos.json')

for d in [MODELOS_DIR, UPLOADS_DIR, os.path.join(BASE_DIR, 'data')]:
    os.makedirs(d, exist_ok=True)

# ─── Persistencia ─────────────────────────────────────────────────────────────

def load_modelos():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_modelos(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ─── Core: modificar docx con python-docx ────────────────────────────────────

def reemplazar_en_parrafo(para, buscar, reemplazar_con):
    """
    Reemplaza texto en un párrafo, incluso si está fragmentado en múltiples runs.
    Une todo el texto, hace el reemplazo y pone el resultado en el primer run.
    """
    texto_completo = ''.join(r.text for r in para.runs)
    if buscar not in texto_completo:
        return False
    nuevo_texto = texto_completo.replace(buscar, reemplazar_con)
    if para.runs:
        para.runs[0].text = nuevo_texto
        for r in para.runs[1:]:
            r.text = ''
    return True

def reemplazar_en_tabla(tabla, buscar, reemplazar_con):
    for fila in tabla.rows:
        for celda in fila.cells:
            for para in celda.paragraphs:
                reemplazar_en_parrafo(para, buscar, reemplazar_con)

def reemplazar_en_doc(doc, buscar, reemplazar_con):
    """Reemplaza en todo el documento: body, headers, footers, tablas."""
    for para in doc.paragraphs:
        reemplazar_en_parrafo(para, buscar, reemplazar_con)
    for tabla in doc.tables:
        reemplazar_en_tabla(tabla, buscar, reemplazar_con)
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr:
                for para in hdr.paragraphs:
                    reemplazar_en_parrafo(para, buscar, reemplazar_con)
                for tabla in hdr.tables:
                    reemplazar_en_tabla(tabla, buscar, reemplazar_con)

def inyectar_firma(doc, firma_path):
    """
    Busca en el documento una imagen placeholder (image2.jpg en las relaciones)
    y la reemplaza con la firma, o inserta la firma sobre la línea de guiones
    del segundo bloque de firma (Franco Gambini).
    Estrategia: reemplaza el archivo image2.jpg dentro del ZIP del docx.
    """
    import zipfile as zf
    # Se hace a nivel ZIP directamente sobre el archivo ya guardado
    # (llamar después de doc.save)
    pass  # lo hacemos en generar_certificado directamente sobre el ZIP

def inyectar_firma_en_zip(docx_path, firma_path):
    """Reemplaza media/image2.jpg dentro del docx con la firma."""
    if not firma_path or not os.path.exists(firma_path):
        return
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in ('word/media/image2.jpg', 'word/media/image2.jpeg'):
                zout.write(firma_path, item.filename)
            else:
                zout.writestr(item, zin.read(item.filename))
    os.replace(tmp, docx_path)

def generar_certificado_docx(template_path, firma_path, piloto, folio_final, fecha_cursada, fecha_validez, output_path):
    """Genera un .docx para un piloto."""
    doc = Document(template_path)

    nombre = piloto['nombre'].upper()
    dni    = piloto['dni']

    # Reemplazar placeholders — soporta múltiples formatos de template
    # Formato nuevo: NOMBRE ALUMNO / NUM_DNI / NUM_FOLIO / FECHA_INICIO / FECHA_FIN
    reemplazar_en_doc(doc, 'NOMBRE ALUMNO', nombre)
    reemplazar_en_doc(doc, 'NUM_DNI', dni)
    reemplazar_en_doc(doc, 'NUM_FOLIO', str(folio_final))
    # Formato anterior: NOMBRE ALUMNO, DNI XXXXX
    reemplazar_en_doc(doc, 'NOMBRE ALUMNO, DNI XXXXX', f'{nombre}, DNI {dni}')
    reemplazar_en_doc(doc, 'DNI XXXXX', f'DNI {dni}')
    if fecha_cursada:
        reemplazar_en_doc(doc, 'FECHA_INICIO', fecha_cursada)
        reemplazar_en_doc(doc, 'DD/MM/YYYY', fecha_cursada)
        reemplazar_en_doc(doc, 'FECHA_CURSADA', fecha_cursada)
    if fecha_validez:
        reemplazar_en_doc(doc, 'FECHA_FIN', fecha_validez)
        reemplazar_en_doc(doc, 'FECHA_VALIDEZ', fecha_validez)
    # Folio en header — formato anterior: XX
    reemplazar_en_doc(doc, 'XX', str(folio_final))

    doc.save(output_path)

    # Inyectar firma reemplazando image2.jpg en el ZIP
    if firma_path and os.path.exists(firma_path):
        inyectar_firma_en_zip(output_path, firma_path)

    return output_path

def docx_a_pdf_windows(docx_path, pdf_dir):
    """Convierte docx a PDF usando LibreOffice (Windows, Linux, cloud)."""
    lo_paths = [
        # Linux / Railway / cloud
        'libreoffice',
        'soffice',
        '/usr/bin/libreoffice',
        '/usr/bin/soffice',
        '/nix/store/libreoffice/bin/soffice',
        # Windows
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        'soffice.exe',
    ]
    soffice = None
    for p in lo_paths:
        try:
            result_test = subprocess.run([p, '--version'], capture_output=True, timeout=5)
            if result_test.returncode == 0:
                soffice = p
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if not soffice:
        return None

    result = subprocess.run(
        [soffice, '--headless', '--convert-to', 'pdf', '--outdir', pdf_dir, docx_path],
        capture_output=True, text=True, timeout=120,
        env={**os.environ, 'HOME': '/tmp'}
    )
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(pdf_dir, base + '.pdf')
    return pdf_path if os.path.exists(pdf_path) else None

# ─── Excel ────────────────────────────────────────────────────────────────────

def leer_excel(path, hoja, col_folio, col_nombre, col_dni):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[hoja] if hoja in wb.sheetnames else wb.active

    ci = int(col_folio) - 1
    cn = int(col_nombre) - 1
    cd = int(col_dni)   - 1

    rows = []
    for row in ws.iter_rows(values_only=True):
        folio  = row[ci] if ci < len(row) else None
        nombre = row[cn] if cn < len(row) else None
        dni    = row[cd] if cd < len(row) else None
        if isinstance(folio, int) and nombre and str(nombre).strip():
            dni_str = str(int(dni)) if isinstance(dni, (float, int)) else str(dni or '').strip()
            rows.append({'folio': folio, 'nombre': str(nombre).strip(), 'dni': dni_str})
    return rows

# ─── Rutas ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html', modelos=load_modelos(), categorias_data=load_categorias())

@app.route('/modelo/nuevo')
def modelo_nuevo():
    cat_default = request.args.get('cat')
    return render_template('modelo_form.html', modelo=None, modelo_id=None, categorias=load_categorias(), cat_default=cat_default)

@app.route('/modelo/<modelo_id>/editar')
def modelo_editar(modelo_id):
    modelos = load_modelos()
    return render_template('modelo_form.html', modelo=modelos.get(modelo_id), modelo_id=modelo_id, categorias=load_categorias(), cat_default=None)

@app.route('/modelo/guardar', methods=['POST'])
def modelo_guardar():
    modelos   = load_modelos()
    modelo_id = request.form.get('modelo_id') or str(uuid.uuid4())[:8]
    nombre    = request.form.get('nombre', '').strip()

    modelo = modelos.get(modelo_id, {'id': modelo_id})
    modelo['nombre'] = nombre
    modelo['categoria_id'] = request.form.get('categoria_id') or None

    if 'docx' in request.files and request.files['docx'].filename:
        f    = request.files['docx']
        dest = os.path.join(MODELOS_DIR, f'{modelo_id}.docx')
        f.save(dest)
        modelo['docx_path'] = dest

    if 'firma' in request.files and request.files['firma'].filename:
        f   = request.files['firma']
        ext = os.path.splitext(f.filename)[1].lower() or '.jpg'
        dest = os.path.join(MODELOS_DIR, f'{modelo_id}_firma{ext}')
        f.save(dest)
        modelo['firma_path'] = dest

    modelos[modelo_id] = modelo
    save_modelos(modelos)
    return redirect(url_for('index'))


@app.route('/modelo/<modelo_id>/descargar-word')
def modelo_descargar_word(modelo_id):
    modelos = load_modelos()
    modelo  = modelos.get(modelo_id)
    if not modelo or not modelo.get('docx_path') or not os.path.exists(modelo['docx_path']):
        return "Archivo no encontrado", 404
    nombre = re.sub(r'[^\w\s-]', '', modelo['nombre']).strip()
    return send_file(modelo['docx_path'], as_attachment=True,
                     download_name=f"{nombre}.docx")

@app.route('/modelo/<modelo_id>/eliminar', methods=['POST'])
def modelo_eliminar(modelo_id):
    modelos = load_modelos()
    if modelo_id in modelos:
        m = modelos.pop(modelo_id)
        for k in ['docx_path', 'firma_path']:
            if p := m.get(k):
                try: os.remove(p)
                except: pass
        save_modelos(modelos)
    return redirect(url_for('index'))

@app.route('/generar')
def generar_page():
    return render_template('generar.html', modelos=load_modelos())

@app.route('/preview-excel', methods=['POST'])
def preview_excel():
    f = request.files.get('excel')
    if not f:
        return jsonify({'error': 'No se recibió archivo'}), 400
    path = os.path.join(UPLOADS_DIR, f'preview_{uuid.uuid4().hex[:6]}.xlsx')
    f.save(path)
    try:
        wb   = openpyxl.load_workbook(path, data_only=True)
        hojas = wb.sheetnames
        ws   = wb.active
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 8: break
            rows.append([str(c) if c is not None else '' for c in list(row)[:12]])
        return jsonify({'hojas': hojas, 'preview': rows, 'tmp_path': path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/preview-hoja', methods=['POST'])
def preview_hoja():
    data = request.json
    try:
        wb = openpyxl.load_workbook(data['path'], data_only=True)
        ws = wb[data['hoja']] if data['hoja'] in wb.sheetnames else wb.active
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 8: break
            rows.append([str(c) if c is not None else '' for c in list(row)[:12]])
        return jsonify({'preview': rows})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generar/ejecutar', methods=['POST'])
def generar_ejecutar():
    data = request.json
    try:
        modelo_id    = data['modelo_id']
        excel_path   = data['excel_path']
        hoja         = data['hoja']
        col_folio    = data['col_folio']
        col_nombre   = data['col_nombre']
        col_dni      = data['col_dni']
        folio_inicio = int(data['folio_inicio'])
        fecha_cursada = data.get('fecha_cursada', '')
        fecha_validez = data.get('fecha_validez', '')

        modelos = load_modelos()
        modelo  = modelos.get(modelo_id)
        if not modelo:
            return jsonify({'error': 'Modelo no encontrado'}), 404
        if not modelo.get('docx_path') or not os.path.exists(modelo['docx_path']):
            return jsonify({'error': 'El modelo no tiene un Word cargado'}), 400

        pilotos = leer_excel(excel_path, hoja, col_folio, col_nombre, col_dni)
        if not pilotos:
            return jsonify({'error': 'No se encontraron filas con datos válidos. Verificá las columnas seleccionadas.'}), 400

        folio_base = pilotos[0]['folio']
        offset     = folio_inicio - folio_base

        tmpbase = os.path.join(UPLOADS_DIR, f'gen_{uuid.uuid4().hex[:8]}')
        os.makedirs(tmpbase)
        pdf_dir = os.path.join(tmpbase, 'pdfs')
        os.makedirs(pdf_dir)

        ok_pdfs = []
        errores = []

        for p in pilotos:
            folio_final  = p['folio'] + offset
            # Nombre del archivo: "Nombre Apellido - Nombre Modelo - Fecha Vencimiento"
            nombre_limpio = re.sub(r'[^\w\s]', '', p['nombre']).strip()
            curso_limpio  = re.sub(r'[^\w\s]', '', modelo['nombre']).strip()
            vto_limpio    = re.sub(r'[/]', '-', fecha_validez) if fecha_validez else 'SIN-VTO'
            nombre_archivo_base = f"{nombre_limpio} - {curso_limpio} - {vto_limpio}"
            nombre_archivo_base = re.sub(r'[\\/:*?"<>|]', '', nombre_archivo_base).strip()[:120]
            docx_out     = os.path.join(tmpbase, f'{nombre_archivo_base}.docx')

            try:
                generar_certificado_docx(
                    modelo['docx_path'],
                    modelo.get('firma_path'),
                    p, folio_final, fecha_cursada, fecha_validez, docx_out
                )
                pdf = docx_a_pdf_windows(docx_out, pdf_dir)
                if pdf:
                    ok_pdfs.append(pdf)
                else:
                    errores.append(f'{p["nombre"]} (PDF fallido)')
            except Exception as e:
                errores.append(f'{p["nombre"]} ({str(e)[:60]})')

        if not ok_pdfs:
            shutil.rmtree(tmpbase, ignore_errors=True)
            return jsonify({'error': f'No se generó ningún PDF. Errores: {errores[:3]}'}), 500

        nom_modelo = re.sub(r'[^\w]', '_', modelo['nombre'])
        zip_name   = f'Certificados_{nom_modelo}_F{folio_inicio}.zip'
        zip_path   = os.path.join(UPLOADS_DIR, zip_name)

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for pdf in sorted(ok_pdfs):
                zf.write(pdf, os.path.basename(pdf))

        shutil.rmtree(tmpbase, ignore_errors=True)

        return jsonify({
            'ok':          len(ok_pdfs),
            'errores':     errores,
            'zip':         zip_name,
            'folio_inicio': folio_inicio,
            'folio_fin':   folio_inicio + len(pilotos) - 1
        })

    except Exception as e:
        import traceback
        return jsonify({'error': traceback.format_exc()}), 500

@app.route('/descargar/<filename>')
def descargar(filename):
    path = os.path.join(UPLOADS_DIR, filename)
    return send_file(path, as_attachment=True, download_name=filename)


@app.route('/actualizar', methods=['POST'])
def actualizar():
    """Auto-actualiza la app via git pull y reinicia Flask."""
    import subprocess, sys, threading, time
    try:
        result = subprocess.run(
            ['git', 'pull', 'origin', 'main'],
            capture_output=True, text=True,
            cwd=BASE_DIR
        )
        if result.returncode == 0:
            output = result.stdout.strip()
            ya_al_dia = 'Already up to date' in output or 'Ya está actualizado' in output
            if not ya_al_dia:
                # Reiniciar Flask en background después de responder
                def reiniciar():
                    time.sleep(1.5)
                    os.execv(sys.executable, [sys.executable] + sys.argv)
                threading.Thread(target=reiniciar, daemon=True).start()
            return jsonify({
                'ok': True,
                'mensaje': 'Ya tenés la última versión.' if ya_al_dia else '¡Actualizado! La app se reinicia sola en 3 segundos...',
                'output': output,
                'actualizado': not ya_al_dia
            })
        else:
            return jsonify({'ok': False, 'mensaje': result.stderr or 'Error al actualizar.'}), 500
    except FileNotFoundError:
        return jsonify({'ok': False, 'mensaje': 'Git no está instalado. Descargalo desde git-scm.com'}), 500
    except Exception as e:
        return jsonify({'ok': False, 'mensaje': str(e)}), 500


# ─── Categorías ───────────────────────────────────────────────────────────────

def load_categorias():
    """Carga categorías. Siempre incluye 'Sin categoría' como fallback."""
    modelos = load_modelos()
    cats_data_file = os.path.join(DATA_ROOT, 'data', 'categorias.json')
    if os.path.exists(cats_data_file):
        with open(cats_data_file, encoding='utf-8') as f:
            cats = json.load(f)
    else:
        cats = []
    return cats

def save_categorias(cats):
    cats_data_file = os.path.join(DATA_ROOT, 'data', 'categorias.json')
    os.makedirs(os.path.dirname(cats_data_file), exist_ok=True)
    with open(cats_data_file, 'w', encoding='utf-8') as f:
        json.dump(cats, f, ensure_ascii=False, indent=2)

@app.route('/categorias')
def categorias_page():
    cats = load_categorias()
    return render_template('categorias.html', categorias=cats)

@app.route('/categoria/guardar', methods=['POST'])
def categoria_guardar():
    data = request.json
    cats = load_categorias()
    nombre = data.get('nombre', '').strip()
    cat_id = data.get('id') or str(uuid.uuid4())[:8]
    color = data.get('color', '#143C64')
    # Actualizar o agregar
    found = False
    for c in cats:
        if c['id'] == cat_id:
            c['nombre'] = nombre
            c['color'] = color
            found = True
            break
    if not found:
        cats.append({'id': cat_id, 'nombre': nombre, 'color': color})
    save_categorias(cats)
    return jsonify({'ok': True, 'id': cat_id})

@app.route('/categoria/<cat_id>/eliminar', methods=['POST'])
def categoria_eliminar(cat_id):
    cats = load_categorias()
    cats = [c for c in cats if c['id'] != cat_id]
    save_categorias(cats)
    # Desasignar modelos que tenían esta categoría
    modelos = load_modelos()
    for m in modelos.values():
        if m.get('categoria_id') == cat_id:
            m['categoria_id'] = None
    save_modelos(modelos)
    return jsonify({'ok': True})

@app.route('/modelo/<modelo_id>/set-categoria', methods=['POST'])
def modelo_set_categoria(modelo_id):
    data = request.json
    modelos = load_modelos()
    if modelo_id in modelos:
        modelos[modelo_id]['categoria_id'] = data.get('categoria_id')
        save_modelos(modelos)
    return jsonify({'ok': True})

if __name__ == '__main__':
    print("\n" + "="*48)
    print("  CertGen JetSMART - Listo!")
    print("  Abri el navegador en: http://localhost:5000")
    print("="*48 + "\n")
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port, threaded=True)
